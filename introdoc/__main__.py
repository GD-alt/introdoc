import pyodbc

from docx import Document
from docx.shared import Pt
from docx.enum.style import WD_STYLE_TYPE

import click

from dataclasses import dataclass
from pathlib import Path


@dataclass
class Column:
    column_name: str
    data_type: str
    is_nullable: bool
    is_primary_key: bool
    is_foreign_key: bool


DATATYPES = {
    -5: 'bigint',
    -2: 'binary',
    -7: 'bit',
    1: 'char',
    91: 'date',
    93: 'datetime',
    -155: 'datetimeoffset',
    3: 'decimal',
    6: 'float',
    -151: 'geography',
    -151: 'geometry',
    -151: 'hierarchyid',
    -4: 'image',
    4: 'int',
    -8: 'nchar',
    -10: 'ntext',
    2: 'numeric',
    -9: 'nvarchar',
    7: 'real',
    5: 'smallint',
    -150: 'sql_variant',
    -1: 'text',
    -154: 'time',
    -2: 'timestamp',
    -6: 'tinyint',
    -11: 'uniqueidentifier',
    -3: 'varbinary',
    12: 'varchar',
    -152: 'xml',
}

size_needed = [
    'binary',
    'char',
    'datetime2',
    'datetimeoffset',
    'decimal',
    'nchar',
    'numeric',
    'nvarchar',
    'time',
    'varbinary',
    'varchar'
]

decimal_needed = [
    'decimal',
    'numeric'
]


def introspect_table(cursor, table) -> list[Column]:
    """
    Introspect a table and return a list of Column objects
    :param cursor: Cursor object
    :param table: Table name
    :return: List of Column objects
    """
    col_list = []

    for col in cursor.columns(table=table):
        if col.type_name in decimal_needed:
            tn = f'{col.type_name}({col.column_size},{col.decimal_digits})'
        elif col.type_name in size_needed:
            tn = f'{col.type_name}({col.column_size})'
        else:
            tn = col.type_name

        col_list.append(
            Column(
                col.column_name,
                tn,
                True if col.is_nullable == 'YES' else False,
                False,
                False
            )
        )

    for col in col_list:
        col.is_primary_key = col.column_name in [key[3] for key in cursor.primaryKeys(table)]
        col.is_foreign_key = col.column_name in [key[3] for key in cursor.foreignKeys(foreignTable=table)]

    return col_list


locales = {
    'en': {
        'primary': 'Primary',
        'foreign': 'Foreign',
        'h_name': 'Column Name',
        'h_datatype': 'Data Type',
        'h_nullable': 'Nullable?',
        'h_keytype': 'Key Type',
        'yes': 'Yes',
        'no': 'No'
    },
    'ru': {
        'primary': 'Первичный',
        'foreign': 'Внешний',
        'h_name': 'Имя колонки',
        'h_datatype': 'Тип данных',
        'h_nullable': 'Принимает NULL?',
        'h_keytype': 'Тип ключа',
        'yes': 'Да',
        'no': 'Нет'
    },
    'de': {
        'primary': 'Primär',
        'foreign': 'Fremd',
        'h_name': 'Spaltenname',
        'h_datatype': 'Datentyp',
        'h_nullable': 'Zulässig NULL?',
        'h_keytype': 'Schlüsseltyp',
        'yes': 'Ja',
        'no': 'Nein'
    }
}


@click.command()
@click.option('--output', '-o', default='output.docx', help='Output file')
@click.option('--database', '-d', default='master', help='Database name')
@click.option('--server', '-s', default='(localdb)\\mssqllocaldb', help='Server name')
@click.option('--driver', '-D', default='ODBC Driver 17 for SQL Server', help='ODBC Driver')
@click.option('--tables', '-t', default=None, help='Tables to introspect')
@click.option('--language', '-l', default='en', help='Language of the document. en, ru, de are supported')
@click.option('--headers', '-h', is_flag=True, help='If to include headers in tables or not')
@click.option('--inullable', '-N', is_flag=True, help='If to include `nullable` column or not')
@click.option('--onatural', '-n', is_flag=True, help='If True, `nullable` column will be `Yes` if column is nullable')
@click.option('--sections', '-s', is_flag=True, help='If to include sections in the document or not')
def main(output, database, server, driver, tables, language, headers, inullable, onatural, sections):
    if language not in ['en', 'ru', 'de']:
        raise ValueError('Language not supported')

    conn = pyodbc.connect(
        f'DRIVER={driver}; SERVER={server}; DATABASE={database};'
    )

    cursor = conn.cursor()

    # Introspect the database
    cursor.execute('SELECT * FROM INFORMATION_SCHEMA.TABLES')

    db_tables = [table[2] for table in cursor.fetchall()]

    tables_dict = {}

    if not tables:
        for table in db_tables:
            tables_dict[table] = introspect_table(cursor, table)

    else:
        tables = [table.strip() for table in tables.split(',')]

        for table in tables:
            if table not in db_tables:
                raise ValueError(f'Table {table} not found in database. Available tables: {", ".join(db_tables)}')

            tables_dict[table] = introspect_table(cursor, table)

    conn.close()

    if Path(output).suffix != '.docx':
        output = f'{output}.docx'

    if Path(output).exists():
        x = input(f'{output} already exists. Overwrite? [y/a/n]: ')

        if x.lower() == 'a':
            doc = Document(output)

        elif x.lower() == 'y':
            Path(output).unlink()
            doc = Document()

        else:
            print('Exiting...')
            return

    else:
        doc = Document()

    styles = doc.styles

    style = styles.add_style('introdoc Section', WD_STYLE_TYPE.PARAGRAPH)

    style.font.size = Pt(14)
    style.font.name = 'Times New Roman'

    style = styles['Table Grid']
    style.font.size = Pt(12)

    for table, columns in tables_dict.items():
        if sections:
            doc.add_paragraph(table).style = 'introdoc Section'

        cols_amnt = 4 if inullable else 3

        if headers:
            table = doc.add_table(rows=1, cols=cols_amnt)
            hdr_cells = table.rows[0].cells
            hdr_cells[0].text = locales[language]['h_name']
            hdr_cells[1].text = locales[language]['h_datatype']
            if inullable:
                hdr_cells[2].text = locales[language]['h_nullable']
                hdr_cells[3].text = locales[language]['h_keytype']
            else:
                hdr_cells[2].text = locales[language]['h_keytype']
        else:
            table = doc.add_table(rows=0, cols=cols_amnt)

        table.style = style

        for col in columns:
            row_cells = table.add_row().cells
            row_cells[0].text = col.column_name
            row_cells[1].text = col.data_type.replace('identity', '')

            if inullable and onatural:
                row_cells[2].text = locales[language]['yes'] if col.is_nullable else locales[language]['no']

            elif inullable and not onatural:
                row_cells[2].text = '+' if col.is_nullable else '-'

            else:
                if col.is_primary_key and col.is_foreign_key and language == 'en':
                    row_cells[2].text = f'{locales[language]["primary"]}, {locales[language]["foreign"]}'
                elif col.is_primary_key and col.is_foreign_key and language in ['ru', 'de']:
                    row_cells[2].text = f'{locales[language]["primary"]}, {locales[language]["foreign"].lower()}'
                elif col.is_primary_key:
                    row_cells[2].text = locales[language]['primary']
                elif col.is_foreign_key:
                    row_cells[2].text = locales[language]['foreign']

            if inullable:
                if col.is_primary_key and col.is_foreign_key and language == 'en':
                    row_cells[3].text = f'{locales[language]["primary"]}, {locales[language]["foreign"]}'
                elif col.is_primary_key and col.is_foreign_key and language in ['ru', 'de']:
                    row_cells[3].text = f'{locales[language]["primary"]}, {locales[language]["foreign"].lower()}'
                elif col.is_primary_key:
                    row_cells[3].text = locales[language]['primary']
                elif col.is_foreign_key:
                    row_cells[3].text = locales[language]['foreign']

    doc.save(output)
    print(f'Document saved to {output}')


if __name__ == '__main__':
    main()
